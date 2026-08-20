from pathlib import Path
from zipfile import ZipFile
from docx import Document
import pdfplumber
from pypdf import PdfReader

docx_path = Path(r"C:\Users\FAJAR\Downloads\25dc836e-ca67-4aeb-a3c1-0894d682626b.docx")
pdf_path = Path(r"C:\Users\FAJAR\Downloads\Laporan-Pelaksanaan-Senam-Tingkat-Kecamatan-Pancoran.pdf")
print("=== DOCX ===")
d = Document(str(docx_path))
for i,p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t:
        print(f"P{i}: {t!r}")
for ti,table in enumerate(d.tables):
    print(f"TABLE {ti}")
    for ri,row in enumerate(table.rows):
        print(f"R{ri}: {[c.text for c in row.cells]!r}")
print("inline_shapes", len(d.inline_shapes))
with ZipFile(docx_path) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print("media", media)
    for n in media:
        print(n, len(z.read(n)))

print("=== PDF ===")
r = PdfReader(str(pdf_path))
print("pages", len(r.pages))
for i,page in enumerate(r.pages):
    print(f"PAGE {i+1} TEXT:")
    print((page.extract_text() or "").strip())
    print("images", len(page.images))
with pdfplumber.open(str(pdf_path)) as pdf:
    for i,p in enumerate(pdf.pages):
        print("pdfplumber page", i+1, "images", len(p.images), "width", p.width, "height", p.height)

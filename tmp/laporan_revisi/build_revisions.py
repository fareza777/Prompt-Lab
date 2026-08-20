from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


BASE = Path(r"C:\Users\FAJAR\Downloads")
WORK = Path(r"E:\apps\promptlab\tmp\laporan_revisi")
SOURCE_DOCX = BASE / "25dc836e-ca67-4aeb-a3c1-0894d682626b.docx"
SOURCE_PDF = BASE / "Laporan-Pelaksanaan-Senam-Tingkat-Kecamatan-Pancoran.pdf"
PHOTO = WORK / "photo1.jpg"
OUT_DOCX = BASE / "25dc836e-ca67-4aeb-a3c1-0894d682626b_perbaikan.docx"
OUT_PDF = BASE / "Laporan-Pelaksanaan-Senam-Tingkat-Kecamatan-Pancoran_perbaikan.pdf"


REPLACEMENTS = {
    "Peserta senam tampak mengenakan kaos olahraga sederhana, didominasi warna putih dan hijau, lengkap dengan celana训练 sebagai pakaian gerak. Mereka memulai dengan pemanasan ringan berupa peregangan otot sebelum masuk ke gerakan inti. Rangkaian gerakan senam dipandu oleh instruktur di bagian depan formasi, sementara peserta mengikuti dengan节奏 secara serempak. Riuh rendah aba-aba instruktur berpadu dengan alunan musik pengiring senam, mengarahkan tempo dan hitungan setiap gerakan agar tetap sinkron. Beberapa peserta terlihat saling membantu menjaga formasi agar barisan tetap rapi di atas paving блок halaman.":
        "Peserta senam tampak mengenakan kaos olahraga sederhana, didominasi warna putih dan hijau, lengkap dengan celana olahraga sebagai pakaian gerak. Mereka memulai dengan pemanasan ringan berupa peregangan otot sebelum masuk ke gerakan inti. Rangkaian gerakan senam dipandu oleh instruktur di bagian depan formasi, sementara peserta mengikuti dengan irama secara serempak. Riuh rendah aba-aba instruktur berpadu dengan alunan musik pengiring senam, mengarahkan tempo dan hitungan setiap gerakan agar tetap sinkron. Beberapa peserta terlihat saling membantu menjaga formasi agar barisan tetap rapi di atas paving block halaman.",
    "Usai senam berakhir, halaman kecamatan tidak langsung kosong. Peserta yang semula berbaris mengikuti senam segera melakukan transisi menuju posisi apel. Bendera merah putih yang terpasang di sisi kanan halaman dan umbul-umbul装饰 lainnya menjadi latar pelaksanaan apel pagi. Apel dipimpin oleh Camat Pancoran, yang berdiri di podium menghadap peserta. Pada sesi ini, Camat Pancoran menyampaikan arahan singkat kepada peserta sebagai bagian dari pembinaan kedisiplinan pegawai dan koordinasi awal pekan.":
        "Usai senam berakhir, halaman kecamatan tidak langsung kosong. Peserta yang semula berbaris mengikuti senam segera melakukan transisi menuju posisi apel. Bendera merah putih yang terpasang di sisi kanan halaman dan umbul-umbul lainnya menjadi latar pelaksanaan apel pagi. Apel dipimpin oleh Camat Pancoran, yang berdiri di podium menghadap peserta. Pada sesi ini, Camat Pancoran menyampaikan arahan singkat kepada peserta sebagai bagian dari pembinaan kedisiplinan pegawai dan koordinasi awal pekan.",
    "Secara keseluruhan, kegiatan berjalan lancar dari awal hingga akhir. Tidak terdapat hambatan berarti, dan peserta mengikuti setiap tahapan kegiatan dengan tertib. Apel pagi diakhiri dengan pembubaran形式 yang tertib sehingga aktivitas pelayanan kecamatan dapat kembali berjalan seperti biasa.":
        "Secara keseluruhan, kegiatan berjalan lancar dari awal hingga akhir. Tidak terdapat hambatan berarti, dan peserta mengikuti setiap tahapan kegiatan dengan tertib. Apel pagi diakhiri dengan pembubaran yang tertib sehingga aktivitas pelayanan kecamatan dapat kembali berjalan seperti biasa.",
    "Dokumentasi kegiatan terdokumentasi pada pukul 07.55 WIB sebagai bukti pelaksanaan kegiatan tingkat kecamatan.":
        "Kegiatan didokumentasikan pada pukul 07.55 WIB sebagai bukti pelaksanaan kegiatan tingkat kecamatan.",
    "Foto 1 — Suasana senam di halaman Kantor Kecamatan Pancoran, peserta mengenakan kaos putih-hijau dan mengikuti gerakan senam secara serempak.":
        "Foto 1 - Suasana senam di halaman Kantor Kecamatan Pancoran, peserta mengenakan kaos putih-hijau dan mengikuti gerakan senam secara serempak.",
}


def set_paragraph_text(paragraph, text):
    """Replace a one-run paragraph while retaining its paragraph and run styling."""
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def add_docx_image_after(document, anchor, image_path):
    image_para = document.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    shape = run.add_picture(str(image_path), width=Inches(5.8))
    shape._inline.docPr.set("descr", "Foto suasana senam di halaman Kantor Kecamatan Pancoran")
    caption = document.add_paragraph("Foto 1")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption.runs:
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].italic = True
    anchor._p.addnext(image_para._p)
    image_para._p.addnext(caption._p)


def build_docx():
    document = Document(str(SOURCE_DOCX))
    paragraphs = list(document.paragraphs)
    for paragraph in paragraphs:
        if paragraph.text in REPLACEMENTS:
            set_paragraph_text(paragraph, REPLACEMENTS[paragraph.text])

    # The source contains no image at all; it also contains a caption for a second
    # photo that was never supplied. Keep the single supplied photo only.
    for paragraph in list(document.paragraphs):
        if paragraph.text.startswith("Foto 2"):
            remove_paragraph(paragraph)

    photo_caption = next(
        (p for p in document.paragraphs if p.text.startswith("Foto 1 -")), None
    )
    if photo_caption is None:
        raise RuntimeError("Foto 1 caption not found in the source DOCX")
    add_docx_image_after(document, photo_caption, PHOTO)
    document.core_properties.title = "Laporan Pelaksanaan Senam Tingkat Kecamatan Pancoran"
    document.core_properties.subject = "Laporan kegiatan senam tingkat kecamatan"
    document.core_properties.author = ""
    document.save(str(OUT_DOCX))


GREEN = colors.HexColor("#27654D")
LIGHT_GREEN = colors.HexColor("#EAF2EE")
GRID = colors.HexColor("#C8D8D0")
TEXT = colors.HexColor("#202824")


def register_fonts():
    arial = Path(r"C:\Windows\Fonts\arial.ttf")
    arial_bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    if arial.exists() and arial_bold.exists():
        pdfmetrics.registerFont(TTFont("Arial", str(arial)))
        pdfmetrics.registerFont(TTFont("Arial-Bold", str(arial_bold)))
        return "Arial", "Arial-Bold"
    return "Helvetica", "Helvetica-Bold"


def draw_page(canvas, doc):
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(GREEN)
    canvas.setLineWidth(0.8)
    canvas.line(25 * mm, 18 * mm, width - 25 * mm, 18 * mm)
    canvas.setFillColor(colors.HexColor("#64716A"))
    canvas.setFont(doc.font_regular, 8.5)
    canvas.drawCentredString(width / 2, 11.5 * mm, f"Halaman {doc.page}")
    canvas.restoreState()


def p(text, style):
    return Paragraph(text, style)


def build_pdf():
    regular, bold = register_fonts()
    left = 25 * mm
    right = 25 * mm
    top = 18 * mm
    bottom = 25 * mm
    width, height = A4
    frame = Frame(left, bottom, width - left - right, height - top - bottom, id="normal")
    doc = BaseDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=left,
        rightMargin=right,
        topMargin=top,
        bottomMargin=bottom,
        title="Laporan Pelaksanaan Senam Tingkat Kecamatan Pancoran",
        author="",
    )
    doc.font_regular = regular
    doc.addPageTemplates([PageTemplate(id="report", frames=frame, onPage=draw_page)])

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], fontName=bold, fontSize=22, leading=25,
        textColor=TEXT, alignment=TA_LEFT, spaceAfter=8,
    )
    date = ParagraphStyle(
        "ReportDate", parent=styles["Normal"], fontName=regular, fontSize=9.5, leading=12,
        textColor=colors.HexColor("#64716A"), spaceAfter=12,
    )
    heading = ParagraphStyle(
        "ReportHeading", parent=styles["Heading2"], fontName=bold, fontSize=13.5, leading=16,
        textColor=TEXT, spaceBefore=7, spaceAfter=7,
    )
    body = ParagraphStyle(
        "ReportBody", parent=styles["BodyText"], fontName=regular, fontSize=9.5, leading=13.2,
        textColor=TEXT, alignment=TA_LEFT, spaceAfter=8,
    )
    bullet = ParagraphStyle(
        "ReportBullet", parent=body, leftIndent=14, firstLineIndent=-9, bulletIndent=0,
        spaceAfter=6,
    )
    table_head = ParagraphStyle(
        "TableHead", parent=body, fontName=bold, textColor=colors.white, spaceAfter=0,
    )
    table_cell = ParagraphStyle(
        "TableCell", parent=body, fontSize=9.2, leading=12, spaceAfter=0,
    )
    photo_caption = ParagraphStyle(
        "PhotoCaption", parent=body, fontSize=9, leading=11, alignment=TA_CENTER,
        textColor=colors.HexColor("#64716A"), spaceBefore=4,
    )

    story = [
        p("Laporan Pelaksanaan Senam Tingkat Kecamatan Pancoran", title),
        p("31 Juli 2026", date),
        Table([[""]], colWidths=[width - left - right], rowHeights=[1.2], style=TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), GREEN),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ])),
        Spacer(1, 6),
        p("Waktu dan Tempat", heading),
    ]

    table_data = [
        [p("Uraian", table_head), p("Keterangan", table_head)],
        [p("Hari/Tanggal", table_cell), p("Jumat, 31 Juli 2026", table_cell)],
        [p("Waktu", table_cell), p("Pagi hari", table_cell)],
        [p("Tempat", table_cell), p("Halaman Kantor Kecamatan Pancoran, Kelurahan Pengadegan, Kecamatan Pancoran, Kota Administrasi Jakarta Selatan", table_cell)],
        [p("Penyelenggara", table_cell), p("Kecamatan Pancoran", table_cell)],
    ]
    info_table = Table(table_data, colWidths=[(width - left - right) * 0.5] * 2, repeatRows=1)
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), GREEN),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, GRID),
        ("BACKGROUND", (0, 1), (-1, 1), colors.white),
        ("BACKGROUND", (0, 2), (-1, 2), LIGHT_GREEN),
        ("BACKGROUND", (0, 3), (-1, 3), colors.white),
        ("BACKGROUND", (0, 4), (-1, 4), LIGHT_GREEN),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(info_table)
    story.extend([
        p("Ringkasan Kegiatan", heading),
        p("Senam tingkat Kecamatan Pancoran dilaksanakan di halaman kantor kecamatan sebagai bagian dari kegiatan pembinaan kebugaran aparatur dan masyarakat di tingkat kecamatan. Setelah sesi senam selesai, kegiatan dilanjutkan dengan apel pagi yang dipimpin langsung oleh Camat Pancoran. Kegiatan didokumentasikan pada pukul 07.55 WIB di lingkungan kantor kecamatan.", body),
        p("Uraian Pelaksanaan", heading),
        p("Senam tingkat Kecamatan Pancoran digelar di halaman Kantor Kecamatan Pancoran, Jakarta Selatan. Sejak pagi, halaman kantor kecamatan yang biasanya digunakan untuk apel sudah ditata dan disiapkan untuk gerakan senam bersama. Lokasi terbuka di depan bangunan utama kecamatan memberikan ruang yang cukup luas bagi peserta untuk bergerak dengan leluasa. Suasana pagi yang cerah turut mendukung berlangsungnya kegiatan secara tertib dan nyaman.", body),
        p("Peserta senam tampak mengenakan kaos olahraga sederhana, didominasi warna putih dan hijau, lengkap dengan celana olahraga sebagai pakaian gerak. Mereka memulai dengan pemanasan ringan berupa peregangan otot sebelum masuk ke gerakan inti. Rangkaian gerakan senam dipandu oleh instruktur di bagian depan formasi, sementara peserta mengikuti dengan irama secara serempak. Riuh rendah aba-aba instruktur berpadu dengan alunan musik pengiring senam, mengarahkan tempo dan hitungan setiap gerakan agar tetap sinkron. Beberapa peserta terlihat saling membantu menjaga formasi agar barisan tetap rapi di atas paving block halaman.", body),
        p("Usai senam berakhir, halaman kecamatan tidak langsung kosong. Peserta yang semula berbaris mengikuti senam segera melakukan transisi menuju posisi apel. Bendera merah putih yang terpasang di sisi kanan halaman dan umbul-umbul lainnya menjadi latar pelaksanaan apel pagi. Apel dipimpin oleh Camat Pancoran, yang berdiri di podium menghadap peserta. Pada sesi ini, Camat Pancoran menyampaikan arahan singkat kepada peserta sebagai bagian dari pembinaan kedisiplinan pegawai dan koordinasi awal pekan.", body),
        p("Secara keseluruhan, kegiatan berjalan lancar dari awal hingga akhir. Tidak terdapat hambatan berarti, dan peserta mengikuti setiap tahapan kegiatan dengan tertib. Apel pagi diakhiri dengan pembubaran yang tertib sehingga aktivitas pelayanan kecamatan dapat kembali berjalan seperti biasa.", body),
        PageBreak(),
        p("Hasil dan Tindak Lanjut", heading),
        p("- Senam Kecamatan Pancoran terlaksana di halaman kantor kecamatan pada Jumat pagi dan diikuti peserta dengan tertib dari pemanasan hingga gerakan inti.", bullet),
        p("- Setelah sesi senam, kegiatan langsung dilanjutkan dengan apel pagi yang dipimpin oleh Camat Pancoran sebagai bagian dari pembinaan kedisiplinan aparatur.", bullet),
        p("- Kegiatan didokumentasikan pada pukul 07.55 WIB sebagai bukti pelaksanaan kegiatan tingkat kecamatan.", bullet),
        p("- Kegiatan senam dan apel pagi rencananya akan menjadi agenda rutin untuk menjaga kebugaran serta memperkuat koordinasi internal kecamatan.", bullet),
        p("- Penanggung jawab kegiatan di tingkat kecamatan agar menyiapkan penjadwalan dan sarana pendukung untuk pelaksanaan senam dan apel pada periode berikutnya.", bullet),
        p("Dokumentasi", heading),
        p("Foto 1 - Suasana senam di halaman Kantor Kecamatan Pancoran, peserta mengenakan kaos putih-hijau dan mengikuti gerakan senam secara serempak.", body),
    ])
    photo = Image(str(PHOTO), width=450, height=337.5)
    photo.hAlign = "CENTER"
    story.extend([photo, p("Foto 1", photo_caption)])
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(OUT_DOCX)
    print(OUT_PDF)

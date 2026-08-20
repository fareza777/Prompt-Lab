from pathlib import Path
from zipfile import ZipFile
from docx import Document
from pypdf import PdfReader

docx=Path(r"C:\Users\FAJAR\Downloads\25dc836e-ca67-4aeb-a3c1-0894d682626b_perbaikan.docx")
pdf=Path(r"C:\Users\FAJAR\Downloads\Laporan-Pelaksanaan-Senam-Tingkat-Kecamatan-Pancoran_perbaikan.pdf")
d=Document(str(docx))
doc_text="\n".join(p.text for p in d.paragraphs)
with ZipFile(docx) as z:
    media=[n for n in z.namelist() if n.startswith("word/media/")]
print("DOCX paragraphs", len(d.paragraphs), "inline_shapes", len(d.inline_shapes), "media", media)
print("DOCX has Foto 2", "Foto 2" in doc_text)
print("DOCX has Chinese", any("\u4e00" <= ch <= "\u9fff" for ch in doc_text))
r=PdfReader(str(pdf))
pdf_text="\n".join((p.extract_text() or "") for p in r.pages)
print("PDF pages", len(r.pages), "page2 images", len(r.pages[1].images))
print("PDF Foto 2 count", pdf_text.count("Foto 2"))
print("PDF has Chinese", any("\u4e00" <= ch <= "\u9fff" for ch in pdf_text))
print("PDF mojibake markers", [m for m in ["celana", "paving block", "umbul-umbul", "pembubaran"] if m not in pdf_text])
